import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def selecionar_arquivo_excel():
    arquivos = [f for f in os.listdir() if f.endswith(('.xls', '.xlsx'))]
    if not arquivos:
        raise FileNotFoundError("Nenhum arquivo Excel encontrado na pasta atual.")
    elif len(arquivos) == 1:
        return arquivos[0]
    else:
        print("Selecione um arquivo:")
        for idx, arquivo in enumerate(arquivos, 1):
            print(f"{idx}. {arquivo}")
        escolha = int(input("Digite o número do arquivo desejado: "))
        return arquivos[escolha - 1]

meses_portugues = {
    'janeiro': 'January', 'fevereiro': 'February', 'março': 'March', 'abril': 'April',
    'maio': 'May', 'junho': 'June', 'julho': 'July', 'agosto': 'August',
    'setembro': 'September', 'outubro': 'October', 'novembro': 'November', 'dezembro': 'December'
}

def traduzir_meses(texto):
    if pd.isna(texto):
        return texto
    for pt, en in meses_portugues.items():
        texto = texto.lower().replace(pt, en)
    return texto

def formatar_data(coluna):
    return pd.to_datetime(
        coluna.apply(traduzir_meses), 
        format='%d %B %Y %H:%M', 
        errors='coerce'
    ).dt.strftime('%d/%m/%y')

def gerar_relatorio(nome_projeto):
    arquivo = selecionar_arquivo_excel()
    df = pd.read_excel(arquivo)

    col_datas = ['Início', 'Término', 'Início_da_Linha_de_Base', 'Término_da_linha_de_base']
    for col in col_datas:
        df[col] = formatar_data(df[col])

    df['Porcentagem_Previsto'] = df['Porcentagem_Previsto'].astype(str).str.replace('%', '').str.replace(',', '.')
    df['Porcentagem_Previsto'] = pd.to_numeric(df['Porcentagem_Previsto'], errors='coerce') / 100

    for col in col_datas:
        df[col + '_DT'] = pd.to_datetime(df[col], format='%d/%m/%y', errors='coerce')
    
    hoje = datetime.now()
    hoje_fmt = hoje.strftime('%d/%m/%y')
    nivel0 = df[df['Nível_da_estrutura_de_tópicos'] == 0].iloc[0]

    AA = nivel0['Término']
    BB = (nivel0['Término_DT'] - nivel0['Término_da_linha_de_base_DT']).days
    CC = nivel0['Término_da_linha_de_base']
    DD = (nivel0['Término_da_linha_de_base_DT'] - nivel0['Início_da_Linha_de_Base_DT']).days
    EE = (nivel0['Término_DT'] - nivel0['Início_DT']).days
    FF = (nivel0['Porcentagem_Concluída'] / nivel0['Porcentagem_Previsto']) if nivel0['Porcentagem_Previsto'] else 0
    FF_fmt = f"{FF:.0%}"

    filtro_horizontes = df[
        df['Nomes_dos_Recursos'].astype(str).str.contains("Horizontes", case=False, na=False) &
        (df['Porcentagem_Concluída'] > 0) &
        (df['Porcentagem_Concluída'] < 1)
    ]

    filtro_cliente = df[
        df['Nomes_dos_Recursos'].astype(str).str.contains("Cliente", case=False, na=False) &
        (df['Porcentagem_Concluída'] > 0) &
        (df['Porcentagem_Concluída'] < 1)
    ]

    def buscar_hierarquia(linha_index):
        pai = avo = bisavo = ''
        for i in range(linha_index - 1, -1, -1):
            nivel = df.at[i, 'Nível_da_estrutura_de_tópicos']
            nome = df.at[i, 'Nome']
            if nivel == 3 and not pai:
                pai = nome
            elif nivel == 2 and not avo:
                avo = nome
            elif nivel == 1 and not bisavo:
                bisavo = nome
            if pai and avo and bisavo:
                break
        return bisavo, avo, pai

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(f"REPORT SEMANAL {nome_projeto.upper()} - {hoje_fmt}")
    run.underline = True
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("📌 RESUMO:")
    resumo1 = f"O projeto, com tendência de término para {AA}, está {BB} dias corridos atrasado em relação à Linha de Base aprovada pelo cliente, que previa término em {CC}."
    resumo2 = f"Com duração inicial de {DD} dias corridos, o projeto possui atualmente duração estimada de {EE} dias corridos."
    resumo3 = f"O grau de aderência do projeto ao planejamento é de {FF_fmt}."

    for texto in [resumo1, resumo2, resumo3]:
        p = doc.add_paragraph(texto)
        p.style = 'List Bullet'

    doc.add_paragraph("\n📅 PRÓXIMAS EMISSÕES DE PROJETO:")
    if filtro_horizontes.empty:
        doc.add_paragraph("- Não existem tarefas que cumpram os critérios desta seção")
    else:
        grupo = {}
        for idx, row in filtro_horizontes.iterrows():
            _, avo, pai = buscar_hierarquia(idx)
            chave = row.get('Subprojeto_Horizontes', 'Não Informado')
            linha = f"{avo} - {pai} - {row['Nome']}: Programado para {row['Término']}"
            grupo.setdefault(chave, []).append(linha)

        for subprojeto, tarefas in grupo.items():
            doc.add_paragraph(f"\n{subprojeto}:")
            for t in tarefas:
                p = doc.add_paragraph(t)
                p.style = 'List Bullet'

    doc.add_paragraph("\n🔎 DISCIPLINAS EM ANÁLISE:")
    if filtro_cliente.empty:
        doc.add_paragraph("- Não existem tarefas que cumpram os critérios desta seção")
    else:
        grupo = {}
        for idx, row in filtro_cliente.iterrows():
            _, avo, pai = buscar_hierarquia(idx)
            chave = row.get('Subprojeto_Horizontes', 'Não Informado')
            dias = (hoje - row['Início_DT']).days if pd.notna(row['Início_DT']) else "?"
            linha = f"{avo} - {pai} - {row['Nome']}: Desde {row['Início']} ({dias} dias)"
            grupo.setdefault(chave, []).append(linha)

        for subprojeto, tarefas in grupo.items():
            doc.add_paragraph(f"\n{subprojeto}:")
            for t in tarefas:
                p = doc.add_paragraph(t)
                p.style = 'List Bullet'

    nome_arquivo = f"Relatorio_Semanal_{nome_projeto.replace(' ', '_')}_{hoje_fmt.replace('/', '-')}.docx"
    doc.save(nome_arquivo)
    print(f"\nRelatório salvo como: {nome_arquivo}")

if __name__ == "__main__":
    projeto = input("Digite o nome do projeto: ")
    gerar_relatorio(projeto)
